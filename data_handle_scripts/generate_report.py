import sys
import pandas as pd
import warnings
import yaml
import os
import json
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')
warnings.filterwarnings('ignore', message='Workbook contains no default style')

# 读取配置
with open('config.yaml', 'r', encoding='utf-8') as f:
    config = yaml.safe_load(f)

start_date = config['date_range']['start_date']
end_date = config['date_range']['end_date']
start_str = start_date.replace('-', '')
end_str = end_date.replace('-', '')


def main():
    # 文件路径：支持 workflow 时间戳
    ts = os.environ.get('WORKFLOW_TIMESTAMP', '')
    suffix = f'_{ts}' if ts else ''
    excel_file = f'中间数据/合同汇总{start_str}-{end_str}{suffix}.xlsx'
    json_file = f'中间数据/统计数据{start_str}-{end_str}{suffix}.json'
    template_file = 'Model/SaaSReportModel.docx'
    output_file = f'结果数据/SaaS月报{start_str}-{end_str}{suffix}.docx'

    # 确保输出目录存在
    os.makedirs('结果数据', exist_ok=True)

    # 读取JSON统计数据
    with open(json_file, 'r', encoding='utf-8') as f:
        stats = json.load(f)

    # 读取Excel数据
    df = pd.read_excel(excel_file)

    # 加载Word模板
    doc = Document(template_file)

    # 替换段落中的占位符
    # 月报月：取结束日期的月份（如6月月报 → 6）
    report_month = int(end_date.split('-')[1])
    placeholder_map = {
        '{{预算使用总额}}': str(stats['预算使用总额']),
        '{{实际运营收益}}': str(stats['实际运营收益']),
        '{{有成本投入专区总数}}': str(stats['有成本投入专区总数']),
        '{{红色}}': str(stats['红色专区个数']),
        '{{橙色}}': str(stats['橙色专区个数']),
        '{{黄色}}': str(stats['黄色专区个数']),
        '{{绿色}}': str(stats['绿色专区个数']),
        '{{月报月}}': str(report_month),
    }

    for para in doc.paragraphs:
        for placeholder, value in placeholder_map.items():
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, value)

    # 填充表格数据
    # 表格1：红色专区
    # 表格2：橙色专区
    # 表格3：黄色专区
    color_sheets = ['红色', '橙色', '黄色']

    for idx, color in enumerate(color_sheets):
        table_idx = idx + 1  # 表格1、2、3
        if table_idx >= len(doc.tables):
            break

        table = doc.tables[table_idx]
        color_df = df[df['颜色分类'] == color].copy()

        if color_df.empty:
            continue

        # 清除现有数据行（保留表头）
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        # 添加数据行
        for i, (_, row) in enumerate(color_df.iterrows(), 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)  # 序号
            row_cells[1].text = str(row.get('分公司', ''))  # 分公司
            row_cells[2].text = str(row.get('合同编号', ''))  # 合同编号
            row_cells[3].text = str(row.get('专区名称', ''))  # 专区名称
            row_cells[4].text = str(row.get('商务', ''))  # 商务
            row_cells[5].text = str(row.get('预算使用', ''))  # 预算使用
            row_cells[6].text = str(row.get('核定总额', ''))  # 核定总额
            row_cells[7].text = str(row.get('预算使用率', ''))  # 预算使用率
            row_cells[8].text = str(row.get('收入成本比', ''))  # 收入成本比
            row_cells[9].text = str(row.get('专属属性', ''))  # 专区属性

    doc.save(output_file)
    print(f"Word报告已生成: {output_file}")

    # 输出统计信息
    print(f"\n填充的占位符:")
    for key, value in placeholder_map.items():
        print(f"  {key} -> {value}")

    print(f"\n表格数据:")
    for color in color_sheets:
        count = len(df[df['颜色分类'] == color])
        print(f"  {color}专区: {count} 行")


if __name__ == '__main__':
    main()
